"""
Generates project_checklist.docx — a professional SDLC checklist
for the CareSkill NGO Learning Platform project.

Run from the project root:
    python generate_checklist.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── colour palette ─────────────────────────────────────────────────────────────
DARK_NAVY   = RGBColor(0x17, 0x32, 0x4D)   # AppColors.ink
BLUE        = RGBColor(0x41, 0xA7, 0xF5)   # AppColors.primary
GREEN       = RGBColor(0x70, 0xD9, 0x8B)   # AppColors.secondary
ORANGE      = RGBColor(0xFF, 0xA2, 0x3A)   # AppColors.accent
SOFT_RED    = RGBColor(0xFF, 0x7D, 0x7D)   # AppColors.softRed
MUTED       = RGBColor(0x6F, 0x7E, 0x8D)   # AppColors.muted
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG    = RGBColor(0xFF, 0xFB, 0xF2)   # AppColors.background
LAVENDER    = RGBColor(0xE9, 0xE2, 0xFF)

# Checkbox characters
CHECK  = "☐"   # unchecked   U+2610
BULLET = "•"

# ── helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    """Set table cell background colour via XML."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_title_page(doc: Document):
    """Decorative cover page."""
    # Top colour strip (table trick)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, DARK_NAVY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(24)
    run = p.add_run("CARESKILL")
    run.bold      = True
    run.font.size = Pt(36)
    run.font.color.rgb = WHITE

    sub = cell.add_paragraph("NGO Learning Platform")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(6)
    r = sub.runs[0] if sub.runs else sub.add_run("NGO Learning Platform")
    r.font.size      = Pt(18)
    r.font.color.rgb = BLUE
    r.bold = True

    tag = cell.add_paragraph("Flutter  ·  FastAPI  ·  SQLite")
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag.paragraph_format.space_after = Pt(24)
    rt = tag.runs[0] if tag.runs else tag.add_run("Flutter  ·  FastAPI  ·  SQLite")
    rt.font.size      = Pt(13)
    rt.font.color.rgb = MUTED

    # Title block below strip
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("PROJECT DEVELOPMENT CHECKLIST")
    tr.bold           = True
    tr.font.size      = Pt(28)
    tr.font.color.rgb = DARK_NAVY

    sub2 = doc.add_paragraph("Complete Software Development Lifecycle")
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s2r = sub2.runs[0] if sub2.runs else sub2.add_run("")
    s2r.font.size      = Pt(14)
    s2r.font.color.rgb = MUTED
    s2r.italic = True

    doc.add_paragraph()

    # Metadata box
    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    rows_data = [
        ("Project",  "CareSkill — NGO Learning Platform"),
        ("Stack",    "Flutter 3 / FastAPI / SQLAlchemy / SQLite"),
        ("Version",  "1.0.0"),
        ("Created",  datetime.date.today().strftime("%d %B %Y")),
    ]
    for i, (label, value) in enumerate(rows_data):
        lc = meta.cell(i, 0)
        vc = meta.cell(i, 1)
        set_cell_bg(lc, DARK_NAVY)
        lc.paragraphs[0].paragraph_format.space_before = Pt(4)
        lc.paragraphs[0].paragraph_format.space_after  = Pt(4)
        lr = lc.paragraphs[0].add_run(label)
        lr.bold           = True
        lr.font.color.rgb = WHITE
        lr.font.size      = Pt(11)

        vr = vc.paragraphs[0].add_run(value)
        vr.font.size      = Pt(11)
        vr.font.color.rgb = DARK_NAVY

    doc.add_paragraph()
    doc.add_paragraph()

    # Legend
    leg = doc.add_paragraph()
    leg.paragraph_format.space_before = Pt(12)
    leg.add_run("Legend:  ").bold = True
    leg.add_run(f"{CHECK} Task to complete   ")
    leg.add_run("★ Optional / Advanced   ")
    leg.add_run("⚠  Security-critical")
    for run in leg.runs:
        run.font.size      = Pt(11)
        run.font.color.rgb = MUTED

    doc.add_page_break()


def add_toc(doc: Document):
    """Simple table-of-contents reference page."""
    h = doc.add_paragraph("TABLE OF CONTENTS", style="Heading 1")
    phases = [
        ("Phase 1",  "Project Planning & Requirements"),
        ("Phase 2",  "Development Environment Setup"),
        ("Phase 3",  "Git & GitHub Workflow"),
        ("Phase 4",  "Database Design & Setup"),
        ("Phase 5",  "Backend Development (Python / FastAPI)"),
        ("Phase 6",  "Frontend Development (Flutter)"),
        ("Phase 7",  "API Integration (Frontend ↔ Backend)"),
        ("Phase 8",  "Testing"),
        ("Phase 9",  "Security"),
        ("Phase 10", "Documentation"),
        ("Phase 11", "Deployment"),
        ("Phase 12", "Post-Deployment & Maintenance"),
        ("Phase 13", "★ Optional Advanced Improvements"),
        ("Phase 14", "User Management & Authentication"),
        ("Phase 15", "Role-Based Access Control (RBAC)"),
        ("Phase 16", "Quiz & Daily Challenge System"),
        ("Phase 17", "Tarase Heerah Drive — Talent Selection"),
        ("Phase 18", "Learning Content Platform"),
    ]
    tbl = doc.add_table(rows=len(phases), cols=2)
    tbl.style = "Table Grid"
    for i, (phase, name) in enumerate(phases):
        lc = tbl.cell(i, 0)
        vc = tbl.cell(i, 1)
        if i % 2 == 0:
            set_cell_bg(lc, RGBColor(0xF0, 0xF4, 0xF8))
            set_cell_bg(vc, RGBColor(0xF0, 0xF4, 0xF8))
        pr = lc.paragraphs[0].add_run(phase)
        pr.bold           = True
        pr.font.color.rgb = BLUE
        pr.font.size      = Pt(11)
        nr = vc.paragraphs[0].add_run(name)
        nr.font.size      = Pt(11)
        nr.font.color.rgb = DARK_NAVY

    doc.add_page_break()


def phase_heading(doc: Document, number: str, title: str, color: RGBColor = BLUE):
    """Coloured phase banner."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(f"  {number}  {title.upper()}")
    run.bold           = True
    run.font.size      = Pt(14)
    run.font.color.rgb = WHITE
    doc.add_paragraph()


def section(doc: Document, title: str):
    """Sub-section heading inside a phase."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f"  {title}")
    run.bold           = True
    run.font.size      = Pt(12)
    run.font.color.rgb = DARK_NAVY
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),  "single")
    bottom.set(qn("w:sz"),   "4")
    bottom.set(qn("w:space"),"1")
    bottom.set(qn("w:color"),"41A7F5")
    pBdr.append(bottom)
    pPr.append(pBdr)


def task(doc: Document, number: str, text: str, indent: int = 0,
         badge: str = "", important: bool = False):
    """Single checklist task row."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Inches(0.3 + indent * 0.25)
    p.paragraph_format.space_before  = Pt(2)
    p.paragraph_format.space_after   = Pt(2)

    cb = p.add_run(f"{CHECK}  ")
    cb.font.size      = Pt(11)
    cb.font.color.rgb = BLUE

    num = p.add_run(f"{number}  ")
    num.bold           = True
    num.font.size      = Pt(11)
    num.font.color.rgb = DARK_NAVY if not important else SOFT_RED

    body = p.add_run(text)
    body.font.size      = Pt(11)
    body.font.color.rgb = DARK_NAVY

    if badge:
        sp = p.add_run(f"  {badge}")
        sp.font.size      = Pt(10)
        sp.font.color.rgb = MUTED
        sp.bold = True


def note(doc: Document, text: str, color: RGBColor = MUTED):
    """Indented grey note line."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.7)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(f"    ↳ {text}")
    r.italic          = True
    r.font.size       = Pt(10)
    r.font.color.rgb  = color


def spacer(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)


# ── CHECKLIST DATA ─────────────────────────────────────────────────────────────

def build_checklist(doc: Document):

    # ── PHASE 1 ────────────────────────────────────────────────────────────────
    phase_heading(doc, "PHASE 1", "Project Planning & Requirements", DARK_NAVY)
    section(doc, "1.1  Project Scope & Goals")
    task(doc, "1.1.1", "Define the core purpose and mission of the application")
    note(doc, "CareSkill: help NGO children aged 8–16 learn skills, track wellness, play games")
    task(doc, "1.1.2", "Identify all user types (roles)")
    note(doc, "Child learner · Parent/Guardian · Counsellor · NGO Admin")
    task(doc, "1.1.3", "Write user stories for each role")
    note(doc, "e.g. 'As a child, I want to log my mood so my counsellor knows how I feel'")
    task(doc, "1.1.4", "Define MVP (Minimum Viable Product) feature list")
    task(doc, "1.1.5", "Separate MVP features from optional/future features")
    task(doc, "1.1.6", "Create a project timeline with milestones and deadlines")
    task(doc, "1.1.7", "Get stakeholder sign-off on scope and timeline")

    section(doc, "1.2  Technology Decisions")
    task(doc, "1.2.1", "Choose frontend framework — Flutter (Dart)")
    task(doc, "1.2.2", "Choose backend framework — FastAPI (Python)")
    task(doc, "1.2.3", "Choose database — SQLite for development, PostgreSQL for production")
    task(doc, "1.2.4", "Choose ORM — SQLAlchemy 2")
    task(doc, "1.2.5", "Choose schema validation — Pydantic v2")
    task(doc, "1.2.6", "Choose HTTP client — http package (Dart)")
    task(doc, "1.2.7", "Document all technology choices and reasons")

    section(doc, "1.3  Project Management Setup")
    task(doc, "1.3.1", "Set up a project management board (GitHub Projects / Trello / Notion)")
    task(doc, "1.3.2", "Create task/issue templates")
    task(doc, "1.3.3", "Define sprint cycle length (e.g. 2-week sprints)")
    task(doc, "1.3.4", "Define coding conventions and naming standards")
    note(doc, "Dart: lowerCamelCase vars, UpperCamelCase classes · Python: snake_case")
    task(doc, "1.3.5", "Define folder structure conventions for frontend and backend")
    task(doc, "1.3.6", "Create a shared communication channel (Slack / WhatsApp / Discord)")
    spacer(doc)

    # ── PHASE 2 ────────────────────────────────────────────────────────────────
    phase_heading(doc, "PHASE 2", "Development Environment Setup", BLUE)
    section(doc, "2.1  Required Tools — Install & Verify")
    task(doc, "2.1.1", "Install Git")
    note(doc, "sudo apt install git  (Linux) · brew install git  (Mac) · git-scm.com  (Windows)")
    task(doc, "2.1.2", "Install Flutter SDK (latest stable)")
    note(doc, "docs.flutter.dev/get-started/install")
    task(doc, "2.1.3", "Install Dart SDK (bundled with Flutter)")
    task(doc, "2.1.4", "Install Android Studio (for Android SDK + emulator)")
    task(doc, "2.1.5", "Install VS Code (recommended editor)")
    task(doc, "2.1.6", "Install Python 3.11")
    note(doc, "sudo apt install python3.11  ·  Check: python3.11 --version")
    task(doc, "2.1.7", "Install uv (fast Python package manager)")
    note(doc, "curl -Ls https://astral.sh/uv/install.sh | sh")
    task(doc, "2.1.8", "Install Postman or Insomnia for API testing")

    section(doc, "2.2  IDE Configuration")
    task(doc, "2.2.1", "Install VS Code Flutter extension (Dart Code)")
    task(doc, "2.2.2", "Install VS Code Python extension (Microsoft)")
    task(doc, "2.2.3", "Install VS Code REST Client extension (for .http files)")
    task(doc, "2.2.4", "Configure VS Code settings (auto format on save, line length 100)")
    task(doc, "2.2.5", "Install Android emulator via Android Studio AVD Manager")
    task(doc, "2.2.6", "Run flutter doctor and resolve all issues")
    note(doc, "flutter doctor  — fix any red ✗ items before continuing")
    task(doc, "2.2.7", "Verify emulator launches: flutter emulators --launch <name>")

    section(doc, "2.3  Python Environment")
    task(doc, "2.3.1", "Create backend folder")
    task(doc, "2.3.2", "Create virtual environment: uv venv .venv --python python3.11")
    task(doc, "2.3.3", "Create requirements.txt with pinned versions")
    note(doc, "fastapi · uvicorn[standard] · sqlalchemy · pydantic · python-docx")
    task(doc, "2.3.4", "Install dependencies: uv pip install -r requirements.txt")
    task(doc, "2.3.5", "Verify installation: .venv/bin/python -c 'import fastapi; print(fastapi.__version__)'")
    spacer(doc)

    # ── PHASE 3 ────────────────────────────────────────────────────────────────
    phase_heading(doc, "PHASE 3", "Git & GitHub Workflow Setup", RGBColor(0x17, 0x32, 0x4D))
    section(doc, "3.1  Local Repository")
    task(doc, "3.1.1", "Navigate to project root folder")
    task(doc, "3.1.2", "Initialise Git repository: git init")
    task(doc, "3.1.3", "Create .gitignore file for Flutter + Python")
    note(doc, "Include: build/ · .dart_tool/ · .venv/ · __pycache__/ · *.db · .env · .DS_Store")
    task(doc, "3.1.4", "Stage all files: git add .")
    task(doc, "3.1.5", "Create initial commit: git commit -m 'chore: initial project setup'")

    section(doc, "3.2  GitHub Remote")
    task(doc, "3.2.1", "Create a new GitHub repository (public or private)")
    task(doc, "3.2.2", "Add remote origin: git remote add origin <url>")
    task(doc, "3.2.3", "Push initial commit: git push -u origin main")
    task(doc, "3.2.4", "Enable branch protection on main (require PR review)")
    task(doc, "3.2.5", "Create develop branch: git checkout -b develop")
    task(doc, "3.2.6", "Push develop branch: git push origin develop")

    section(doc, "3.3  Branching & Commit Conventions")
    task(doc, "3.3.1", "Define branch naming: feature/*, bugfix/*, hotfix/*, chore/*")
    task(doc, "3.3.2", "Define commit message format (Conventional Commits)")
    note(doc, "feat: · fix: · docs: · style: · refactor: · test: · chore:")
    task(doc, "3.3.3", "Write CONTRIBUTING.md with Git workflow guide")
    task(doc, "3.3.4", "Set up GitHub Issue templates (Bug Report, Feature Request)")
    task(doc, "3.3.5", "Set up GitHub Pull Request template")
    task(doc, "3.3.6", "Create GitHub Actions workflow for CI (run flutter test + pytest on PR)")
    spacer(doc)

    # ── PHASE 4 ────────────────────────────────────────────────────────────────
    phase_heading(doc, "PHASE 4", "Database Design & Setup", GREEN)
    section(doc, "4.1  Schema Design")
    task(doc, "4.1.1", "Draw Entity-Relationship (ER) diagram")
    note(doc, "Tools: dbdiagram.io · draw.io · Mermaid erDiagram in README")
    task(doc, "4.1.2", "Define all tables: users, skill_categories, courses, user_course_progress")
    task(doc, "4.1.3", "Define tables: games, game_sessions, counselling_sessions")
    task(doc, "4.1.4", "Define tables: mood_logs, badges, user_badges")
    task(doc, "4.1.5", "Define all primary keys (id INTEGER, autoincrement)")
    task(doc, "4.1.6", "Define all foreign keys and cascade rules (ON DELETE CASCADE)")
    task(doc, "4.1.7", "Define column types, nullable constraints, and defaults")
    task(doc, "4.1.8", "Add indexes on frequently queried columns (user_id FKs)")

    section(doc, "4.2  SQLAlchemy ORM Setup")
    task(doc, "4.2.1", "Create database.py with engine, SessionLocal, and Base")
    task(doc, "4.2.2", "Create models/user.py — User model with all relationships")
    task(doc, "4.2.3", "Create models/course.py — SkillCategory, Course, UserCourseProgress")
    task(doc, "4.2.4", "Create models/game.py — Game, GameSession")
    task(doc, "4.2.5", "Create models/wellness.py — CounsellingSession, MoodLog")
    task(doc, "4.2.6", "Create models/badge.py — Badge, UserBadge")
    task(doc, "4.2.7", "Import all models in models/__init__.py so Base.metadata finds them")
    task(doc, "4.2.8", "Run Base.metadata.create_all(engine) and verify tables are created")

    section(doc, "4.3  Seed Data")
    task(doc, "4.3.1", "Create seed.py script")
    task(doc, "4.3.2", "Seed skill categories (Coding, Cyber Safety, Communication, Art, Music, Languages)")
    task(doc, "4.3.3", "Seed courses (Coding Basics, Speak with Confidence, Internet Safety Heroes)")
    task(doc, "4.3.4", "Seed games (Quiz Battle, Memory Game, Puzzle Game, Safety Adventure, etc.)")
    task(doc, "4.3.5", "Seed badges (Coding Starter, Cyber Safe, Speaker, Creative Mind, etc.)")
    task(doc, "4.3.6", "Seed demo user: Aarav Sharma, age 12, level 4, 2480 XP")
    task(doc, "4.3.7", "Seed user course progress, game sessions, counselling session, mood logs")
    task(doc, "4.3.8", "Run seed.py and verify data in SQLite using DB Browser or sqlite3 CLI")
    spacer(doc)

    # ── PHASE 5 ────────────────────────────────────────────────────────────────
    phase_heading(doc, "PHASE 5", "Backend Development — Python / FastAPI", ORANGE)
    section(doc, "5.1  Project Structure")
    task(doc, "5.1.1", "Create backend/ folder structure: app/models, schemas, crud, routers")
    task(doc, "5.1.2", "Create app/config.py with Settings class (DB URL, XP per level, debug flag)")
    task(doc, "5.1.3", "Create app/database.py with engine, SessionLocal, get_db dependency")
    task(doc, "5.1.4", "Create app/main.py with FastAPI() instance and CORS middleware")

    section(doc, "5.2  Pydantic Schemas")
    task(doc, "5.2.1", "Create schemas/user.py — UserCreate, UserUpdate, XPAdd, UserResponse, UserStats, LeaderboardEntry")
    task(doc, "5.2.2", "Create schemas/course.py — CategoryResponse, CourseResponse, ProgressUpdate, UserCourseProgressResponse")
    task(doc, "5.2.3", "Create schemas/game.py — GameResponse, GameSessionCreate, GameSessionResponse")
    task(doc, "5.2.4", "Create schemas/wellness.py — CounsellingCreate/Update/Response, MoodCreate/Response, enums")
    task(doc, "5.2.5", "Create schemas/badge.py — BadgeResponse, UserBadgeResponse")
    task(doc, "5.2.6", "Validate all schemas parse correctly with test data")

    section(doc, "5.3  CRUD Layer")
    task(doc, "5.3.1", "Create crud/user_crud.py — get_user, get_users, create_user, update_user, add_xp, get_user_stats")
    note(doc, "add_xp must auto-recalculate level: level = xp // xp_per_level + 1")
    task(doc, "5.3.2", "Create crud/course_crud.py — get_categories, get_courses, get_course, get_user_course_progress, upsert_course_progress")
    task(doc, "5.3.3", "Create crud/game_crud.py — get_games, get_game, record_game_session, get_user_game_history")
    task(doc, "5.3.4", "Create crud/wellness_crud.py — get/create/update counselling sessions, log/get mood")
    task(doc, "5.3.5", "Create crud/badge_crud.py — get_badges, get_user_badges, award_badge (idempotent)")

    section(doc, "5.4  API Routers")
    task(doc, "5.4.1", "Create routers/users.py — GET/POST /users/, GET/PATCH /users/{id}, POST /users/{id}/xp, GET /users/{id}/stats")
    task(doc, "5.4.2", "Create routers/courses.py — GET /categories, GET /courses, GET /courses/{id}, GET+PUT /users/{id}/courses")
    task(doc, "5.4.3", "Create routers/games.py — GET /games, POST /users/{id}/games/{gid}/sessions, GET game history")
    note(doc, "POST session must call add_xp to award XP to the user automatically")
    task(doc, "5.4.4", "Create routers/wellness.py — GET/POST counselling, PATCH session, POST/GET mood")
    task(doc, "5.4.5", "Create routers/badges.py — GET /badges, GET/POST /users/{id}/badges")
    task(doc, "5.4.6", "Create routers/leaderboard.py — GET /leaderboard/, GET /leaderboard/{id}/rank")
    task(doc, "5.4.7", "Register all routers in main.py with app.include_router()")

    section(doc, "5.5  Endpoint Verification")
    task(doc, "5.5.1", "Start server: uvicorn app.main:app --reload --port 8000")
    task(doc, "5.5.2", "Open Swagger UI at http://localhost:8000/docs")
    task(doc, "5.5.3", "Test every endpoint via Swagger UI (try 200 and 404 responses)")
    task(doc, "5.5.4", "Test all endpoints via Postman — export collection for team use")
    task(doc, "5.5.5", "Verify 404 responses for missing IDs")
    task(doc, "5.5.6", "Verify 422 responses for invalid request bodies")
    task(doc, "5.5.7", "Verify GET /leaderboard returns users sorted by XP descending")
    spacer(doc)

    # ── PHASE 6 ────────────────────────────────────────────────────────────────
    phase_heading(doc, "PHASE 6", "Frontend Development — Flutter", BLUE)
    section(doc, "6.1  Project Bootstrap")
    task(doc, "6.1.1", "Run: flutter create flutter_application_1 (or initialise existing project)")
    task(doc, "6.1.2", "Run: flutter pub get to install dependencies")
    task(doc, "6.1.3", "Create folder structure: lib/core · lib/models · lib/services · lib/utils · lib/widgets · lib/screens")
    task(doc, "6.1.4", "Add http: ^1.2.0 to pubspec.yaml dependencies")
    task(doc, "6.1.5", "Run flutter pub get again after adding http")

    section(doc, "6.2  Core & Shared Files")
    task(doc, "6.2.1", "Create lib/core/colors.dart — AppColors class with all named colours")
    task(doc, "6.2.2", "Create lib/app_state.dart — AppState with static userId constant")
    task(doc, "6.2.3", "Create lib/utils/icon_mapper.dart — maps icon name strings to IconData")
    task(doc, "6.2.4", "Configure MaterialApp in main.dart with Material 3 theme and custom font")
    task(doc, "6.2.5", "Create AppShell StatefulWidget with bottom NavigationBar (5 tabs)")

    section(doc, "6.3  Shared Widgets")
    task(doc, "6.3.1", "Create widgets/app_card.dart — Card container with rounded corners + shadow")
    task(doc, "6.3.2", "Create widgets/app_scroll_view.dart — ListView wrapper with consistent padding")
    task(doc, "6.3.3", "Create widgets/top_header.dart — Title, subtitle, action icon button")
    task(doc, "6.3.4", "Create widgets/section_header.dart — Section label + optional action text button")
    task(doc, "6.3.5", "Create widgets/search_box.dart — Filled text field with search icon")
    task(doc, "6.3.6", "Create widgets/filter_chip_label.dart — Chip with optional selected state")

    section(doc, "6.4  Data Models")
    task(doc, "6.4.1", "Create models/skill_category.dart with id, fromJson factory, const data list")
    task(doc, "6.4.2", "Create models/course.dart with id, fromJson, fromProgressJson factories")
    task(doc, "6.4.3", "Create models/game_item.dart with id, fromJson factory, const data list")
    task(doc, "6.4.4", "Create models/api_models.dart — AppUser, UserStats, LeaderboardEntry, ApiCounsellingSession, ApiBadge, UserBadge")

    section(doc, "6.5  Screens")
    task(doc, "6.5.1", "Implement HomeDashboard — skill categories row, daily challenge, safety story, counselling card")
    task(doc, "6.5.2", "Implement LearnScreen — search box, filter chips, course cards with progress bars")
    task(doc, "6.5.3", "Implement GamesScreen — leaderboard card, 2-column game grid, admin invite card")
    task(doc, "6.5.4", "Implement WellnessScreen — emergency card, action grid, mood tracker, booking flow, protection lesson")
    task(doc, "6.5.5", "Implement ProfileScreen — user hero card, badges wrap, analytics grid, counselling history")
    task(doc, "6.5.6", "Verify all screens render correctly with static/hardcoded data first")
    task(doc, "6.5.7", "Test UI on small screen (360dp) and large screen (412dp)")
    task(doc, "6.5.8", "Test portrait and landscape orientations")
    spacer(doc)

    # ── PHASE 7 ────────────────────────────────────────────────────────────────
    phase_heading(doc, "PHASE 7", "API Integration — Frontend ↔ Backend", ORANGE)
    section(doc, "7.1  HTTP Client Setup")
    task(doc, "7.1.1", "Create lib/services/api_client.dart with static baseUrl constant")
    task(doc, "7.1.2", "Implement ApiClient.get() with JSON decoding and timeout")
    task(doc, "7.1.3", "Implement ApiClient.post(), put(), patch() methods")
    task(doc, "7.1.4", "Create ApiException class for non-2xx status codes")
    task(doc, "7.1.5", "Add INTERNET permission to android/app/src/main/AndroidManifest.xml", important=True)
    task(doc, "7.1.6", "Add android:usesCleartextTraffic='true' to application tag (dev only)", important=True)

    section(doc, "7.2  Service Classes")
    task(doc, "7.2.1", "Create services/user_service.dart — getUser, getUserStats, addXp")
    task(doc, "7.2.2", "Create services/course_service.dart — getCategories, getCourses, getUserCourses, updateProgress")
    task(doc, "7.2.3", "Create services/game_service.dart — getGames, recordSession")
    task(doc, "7.2.4", "Create services/wellness_service.dart — getCounsellingSessions, bookSession, logMood")
    task(doc, "7.2.5", "Create services/badge_service.dart — getUserBadges, awardBadge")
    task(doc, "7.2.6", "Create services/leaderboard_service.dart — getLeaderboard, getUserRank")

    section(doc, "7.3  Screen Wiring")
    task(doc, "7.3.1", "Convert HomeDashboard to StatefulWidget — load categories + upcoming session")
    task(doc, "7.3.2", "Convert LearnScreen to StatefulWidget — load getUserCourses with real progress")
    task(doc, "7.3.3", "Convert GamesScreen to StatefulWidget — load games list from API")
    task(doc, "7.3.4", "Convert WellnessScreen MoodTrackerCard to StatefulWidget — POST mood on tap")
    task(doc, "7.3.5", "Convert ProfileScreen to StatefulWidget — load user, stats, and badges")
    task(doc, "7.3.6", "Update CounsellingSessionCard to accept ApiCounsellingSession? parameter")
    task(doc, "7.3.7", "Update ProfileHeroCard to accept AppUser? and display real data")

    section(doc, "7.4  UX: Loading & Error States")
    task(doc, "7.4.1", "Show CircularProgressIndicator while API call is in-flight")
    task(doc, "7.4.2", "Show error view with descriptive message on network failure")
    task(doc, "7.4.3", "Add Retry button in all error views")
    task(doc, "7.4.4", "Check if mounted before calling setState() after async operations")
    task(doc, "7.4.5", "Verify 'Mood logged ✓' confirmation appears after mood tap")
    task(doc, "7.4.6", "Verify counselling card shows 'No upcoming sessions' when empty")

    section(doc, "7.5  End-to-End Verification")
    task(doc, "7.5.1", "Start backend server and run Flutter app simultaneously")
    task(doc, "7.5.2", "Verify Home screen loads skill categories from /categories endpoint")
    task(doc, "7.5.3", "Verify Learn screen shows real course progress bars from /users/1/courses")
    task(doc, "7.5.4", "Verify Games screen loads 6 games from /games endpoint")
    task(doc, "7.5.5", "Verify mood tap sends POST to /users/1/wellness/mood and shows confirmation")
    task(doc, "7.5.6", "Verify Profile shows Aarav Sharma with 2480 XP, level 4")
    task(doc, "7.5.7", "Verify profile badges loaded from /users/1/badges")
    task(doc, "7.5.8", "Verify analytics tiles show data from /users/1/stats")
    spacer(doc)

    # ── PHASE 8 ────────────────────────────────────────────────────────────────
    phase_heading(doc, "PHASE 8", "Testing", GREEN)
    section(doc, "8.1  Backend Unit Tests (pytest)")
    task(doc, "8.1.1", "Install pytest and httpx: uv pip install pytest httpx")
    task(doc, "8.1.2", "Create tests/ folder in backend with conftest.py (test DB fixture)")
    task(doc, "8.1.3", "Write unit tests for user_crud — create, get, update, add_xp")
    task(doc, "8.1.4", "Write unit tests for level-up logic (xp // xp_per_level + 1)")
    task(doc, "8.1.5", "Write unit tests for course_crud — upsert_course_progress, completed flag")
    task(doc, "8.1.6", "Write unit tests for badge_crud — award_badge idempotency")
    task(doc, "8.1.7", "Write unit tests for wellness_crud — mood log, counselling session status")
    task(doc, "8.1.8", "Write unit tests for leaderboard rank calculation")

    section(doc, "8.2  Backend Integration Tests (FastAPI TestClient)")
    task(doc, "8.2.1", "Test POST /users/ creates user and returns 201")
    task(doc, "8.2.2", "Test GET /users/{id} returns correct user data")
    task(doc, "8.2.3", "Test GET /users/999 returns 404")
    task(doc, "8.2.4", "Test POST /users/{id}/xp adds XP and updates level")
    task(doc, "8.2.5", "Test PUT /users/{id}/courses/{cid}/progress sets progress and completed flag")
    task(doc, "8.2.6", "Test POST /users/{id}/games/{gid}/sessions creates session and awards XP")
    task(doc, "8.2.7", "Test POST /users/{id}/wellness/mood stores mood correctly")
    task(doc, "8.2.8", "Test GET /leaderboard/ returns users sorted by XP descending")
    task(doc, "8.2.9", "Test Pydantic schema validation rejects invalid mood values (422)")
    task(doc, "8.2.10", "Run full suite: .venv/bin/pytest -v")

    section(doc, "8.3  Flutter Widget Tests")
    task(doc, "8.3.1", "Write widget test for AppCard renders child with correct padding")
    task(doc, "8.3.2", "Write widget test for TopHeader renders title and subtitle")
    task(doc, "8.3.3", "Write widget test for CourseCard renders title, level, and progress bar")
    task(doc, "8.3.4", "Write widget test for MoodTrackerCard shows 5 mood buttons")
    task(doc, "8.3.5", "Write widget test for ProfileHeroCard shows name/XP from AppUser")
    task(doc, "8.3.6", "Write widget test for loading state (CircularProgressIndicator visible)")
    task(doc, "8.3.7", "Write widget test for error state (error message + Retry button visible)")
    task(doc, "8.3.8", "Run: flutter test")

    section(doc, "8.4  Manual & Exploratory Testing")
    task(doc, "8.4.1", "Test all 5 navigation tabs switch screens correctly")
    task(doc, "8.4.2", "Test app on Android emulator API 34 (Pixel 6 profile)")
    task(doc, "8.4.3", "Test app on small device (Pixel 3a, 5.6-inch screen)")
    task(doc, "8.4.4", "Test network failure scenario (stop backend, verify error view appears)")
    task(doc, "8.4.5", "Test with empty database (fresh .db file, no seed data)")
    task(doc, "8.4.6", "Test mood tracker rapid-tap (should not double-submit)")
    task(doc, "8.4.7", "Test Retry button reloads data successfully after server restart")
    spacer(doc)

    # ── PHASE 9 ────────────────────────────────────────────────────────────────
    phase_heading(doc, "PHASE 9", "Security", SOFT_RED)
    section(doc, "9.1  Backend Security")
    task(doc, "9.1.1", "Validate all incoming request body fields with Pydantic", important=True)
    task(doc, "9.1.2", "Use parameterised SQLAlchemy queries (never raw string SQL)", important=True)
    note(doc, "SQLAlchemy ORM queries are parameterised by default — never use f-strings in queries")
    task(doc, "9.1.3", "Restrict CORS to specific origins in production (not allow_origins=['*'])", important=True)
    task(doc, "9.1.4", "Add API rate limiting (slowapi library)")
    task(doc, "9.1.5", "Move all secrets and config to .env file (use python-dotenv)")
    task(doc, "9.1.6", "Add .env to .gitignore — never commit secrets to Git", important=True)
    task(doc, "9.1.7", "Add authentication: implement JWT-based login endpoint", important=True)
    note(doc, "Use python-jose and passlib — return JWT token on successful login")
    task(doc, "9.1.8", "Add authentication middleware to protect all user-specific endpoints")
    task(doc, "9.1.9", "Validate that users can only access their own data (not /users/2/badges as user 1)")
    task(doc, "9.1.10", "Log API errors and suspicious requests to a log file")
    task(doc, "9.1.11", "Add request size limits to prevent large payload attacks")

    section(doc, "9.2  Frontend Security")
    task(doc, "9.2.1", "Remove android:usesCleartextTraffic='true' before production build", important=True)
    task(doc, "9.2.2", "Use HTTPS for all production API calls", important=True)
    task(doc, "9.2.3", "Store JWT token in flutter_secure_storage (not SharedPreferences)", important=True)
    task(doc, "9.2.4", "Never hardcode API keys or secrets in Dart code", important=True)
    task(doc, "9.2.5", "Validate all user inputs before sending to API")
    task(doc, "9.2.6", "Add certificate pinning for production mobile app (advanced)")
    task(doc, "9.2.7", "Remove all debug print() / debugPrint() statements before production build")
    task(doc, "9.2.8", "Enable Flutter obfuscation for release build: --obfuscate --split-debug-info")

    section(doc, "9.3  Child Safety (NGO-specific)")
    task(doc, "9.3.1", "Ensure all counselling session data is accessible only to authorised users", important=True)
    task(doc, "9.3.2", "Ensure mood logs are private (child + counsellor only)", important=True)
    task(doc, "9.3.3", "Implement parent PIN or password before accessing Parent Dashboard")
    task(doc, "9.3.4", "Add content moderation for any user-generated text")
    task(doc, "9.3.5", "Ensure emergency help button works without authentication")
    spacer(doc)

    # ── PHASE 10 ───────────────────────────────────────────────────────────────
    phase_heading(doc, "PHASE 10", "Documentation", DARK_NAVY)
    section(doc, "10.1  Project README")
    task(doc, "10.1.1", "Write project overview and mission statement")
    task(doc, "10.1.2", "Add feature list table")
    task(doc, "10.1.3", "Add tech stack table")
    task(doc, "10.1.4", "Add system architecture Mermaid diagram")
    task(doc, "10.1.5", "Add app navigation flowchart (screen → widgets → API)")
    task(doc, "10.1.6", "Add database ER diagram")
    task(doc, "10.1.7", "Add API request-response sequence diagram")
    task(doc, "10.1.8", "Add backend layer architecture diagram")
    task(doc, "10.1.9", "Add step-by-step Quick Start guide (backend + Flutter)")
    task(doc, "10.1.10", "Add Android emulator / iOS / physical device URL note")
    task(doc, "10.1.11", "Add complete API endpoint reference table")
    task(doc, "10.1.12", "Add configuration reference section")

    section(doc, "10.2  Code Documentation")
    task(doc, "10.2.1", "Add docstrings to all Python functions (models, CRUD, routers)")
    task(doc, "10.2.2", "Add /// doc comments to public Dart classes and methods")
    task(doc, "10.2.3", "Document non-obvious logic with inline comments")
    note(doc, "e.g. the level-up formula, idempotent badge award, mounted check after async")
    task(doc, "10.2.4", "Add OpenAPI descriptions to FastAPI endpoints (description= parameter)")
    task(doc, "10.2.5", "Ensure Swagger UI at /docs is accurate and helpful")

    section(doc, "10.3  Team Documentation")
    task(doc, "10.3.1", "Create this project checklist (project_checklist.docx)")
    task(doc, "10.3.2", "Write CONTRIBUTING.md (Git workflow, code style guide, PR process)")
    task(doc, "10.3.3", "Write CHANGELOG.md (track features added per version)")
    task(doc, "10.3.4", "Create Postman collection and export as JSON")
    task(doc, "10.3.5", "Write database migration guide (for future schema changes)")
    spacer(doc)

    # ── PHASE 11 ───────────────────────────────────────────────────────────────
    phase_heading(doc, "PHASE 11", "Deployment", ORANGE)
    section(doc, "11.1  Backend Deployment (Cloud)")
    task(doc, "11.1.1", "Choose cloud hosting: Railway / Render / Fly.io / DigitalOcean")
    task(doc, "11.1.2", "Migrate database from SQLite → PostgreSQL for production", important=True)
    note(doc, "Install psycopg2-binary · change database_url in config · test migrations")
    task(doc, "11.1.3", "Create .env file for production secrets (DATABASE_URL, SECRET_KEY)")
    task(doc, "11.1.4", "Create Procfile or start command: uvicorn app.main:app --host 0.0.0.0 --port $PORT")
    task(doc, "11.1.5", "Configure environment variables in the hosting dashboard")
    task(doc, "11.1.6", "Enable HTTPS (most cloud platforms provide this automatically)")
    task(doc, "11.1.7", "Deploy and verify server health: GET https://your-api.com/")
    task(doc, "11.1.8", "Run production seed or migration script")
    task(doc, "11.1.9", "Test all production API endpoints via Postman")
    task(doc, "11.1.10", "Restrict CORS to production Flutter app origin", important=True)

    section(doc, "11.2  Flutter Android Release Build")
    task(doc, "11.2.1", "Update baseUrl in api_client.dart to production HTTPS URL", important=True)
    task(doc, "11.2.2", "Remove usesCleartextTraffic='true' from AndroidManifest.xml", important=True)
    task(doc, "11.2.3", "Update app version in pubspec.yaml (version: 1.0.0+1)")
    task(doc, "11.2.4", "Generate keystore: keytool -genkeypair -v -keystore release.jks ...")
    task(doc, "11.2.5", "Configure signing in android/app/build.gradle")
    task(doc, "11.2.6", "Add keystore file path and passwords to key.properties (add to .gitignore)")
    task(doc, "11.2.7", "Build release APK: flutter build apk --release --obfuscate --split-debug-info=symbols/")
    task(doc, "11.2.8", "Test release APK on a real Android device")
    task(doc, "11.2.9", "Build App Bundle: flutter build appbundle --release")
    task(doc, "11.2.10", "Upload to Google Play Console (internal testing track)")

    section(doc, "11.3  CI/CD Pipeline")
    task(doc, "11.3.1", "Create .github/workflows/ci.yml")
    task(doc, "11.3.2", "Add job: flutter analyze + flutter test on every push")
    task(doc, "11.3.3", "Add job: pytest on every push")
    task(doc, "11.3.4", "Add job: auto-deploy backend on merge to main")
    task(doc, "11.3.5", "Add status badges to README (build passing / test passing)")
    spacer(doc)

    # ── PHASE 12 ───────────────────────────────────────────────────────────────
    phase_heading(doc, "PHASE 12", "Post-Deployment & Maintenance", GREEN)
    section(doc, "12.1  Monitoring & Observability")
    task(doc, "12.1.1", "Set up error tracking — Sentry for both FastAPI and Flutter")
    task(doc, "12.1.2", "Monitor API response times and error rates")
    task(doc, "12.1.3", "Set up uptime monitoring (UptimeRobot or Better Uptime)")
    task(doc, "12.1.4", "Review API logs weekly for unusual patterns")
    task(doc, "12.1.5", "Set up database query performance alerts")

    section(doc, "12.2  Data & Backups")
    task(doc, "12.2.1", "Schedule daily automated database backups")
    task(doc, "12.2.2", "Test backup restore process")
    task(doc, "12.2.3", "Store backups in a separate location (S3 / Google Drive)")
    task(doc, "12.2.4", "Define data retention policy for mood logs and session notes")

    section(doc, "12.3  Dependency Maintenance")
    task(doc, "12.3.1", "Run flutter pub outdated monthly — update compatible packages")
    task(doc, "12.3.2", "Run uv pip list --outdated monthly — update Python packages")
    task(doc, "12.3.3", "Review GitHub Dependabot alerts for security vulnerabilities")
    task(doc, "12.3.4", "Test app after every dependency update")

    section(doc, "12.4  Continuous Improvement")
    task(doc, "12.4.1", "Collect user feedback from children, counsellors, and NGO staff")
    task(doc, "12.4.2", "Analyse usage analytics to identify most-used features")
    task(doc, "12.4.3", "Plan next sprint based on feedback and analytics")
    task(doc, "12.4.4", "Review and update CHANGELOG.md for each release")
    task(doc, "12.4.5", "Conduct quarterly security review")
    spacer(doc)

    # ── PHASE 13 ───────────────────────────────────────────────────────────────
    phase_heading(doc, "PHASE 13", "★ Optional Advanced Improvements", MUTED)
    section(doc, "13.1  Authentication & User Management")
    task(doc, "13.1.1", "Add JWT-based login/register endpoints", badge="★ Advanced")
    task(doc, "13.1.2", "Add Flutter login screen with email + password", badge="★ Advanced")
    task(doc, "13.1.3", "Add role-based access control (child / counsellor / admin)", badge="★ Advanced")
    task(doc, "13.1.4", "Add parent account with PIN-protected child view", badge="★ Advanced")

    section(doc, "13.2  Real-Time Features")
    task(doc, "13.2.1", "Add WebSocket support to FastAPI for real-time quiz battle", badge="★ Advanced")
    task(doc, "13.2.2", "Implement live leaderboard updates via WebSocket", badge="★ Advanced")
    task(doc, "13.2.3", "Add push notifications with Firebase FCM (session reminders)", badge="★ Advanced")
    task(doc, "13.2.4", "Add in-app chat between child and counsellor", badge="★ Advanced")

    section(doc, "13.3  Offline Mode")
    task(doc, "13.3.1", "Add sqflite to Flutter for local SQLite caching", badge="★ Advanced")
    task(doc, "13.3.2", "Cache courses and games for offline browsing", badge="★ Advanced")
    task(doc, "13.3.3", "Queue mood logs and sync when online", badge="★ Advanced")
    task(doc, "13.3.4", "Add connectivity_plus to detect network state", badge="★ Advanced")

    section(doc, "13.4  Multimedia & Content")
    task(doc, "13.4.1", "Integrate video_player for course video content", badge="★ Advanced")
    task(doc, "13.4.2", "Add audio narration for safety stories", badge="★ Advanced")
    task(doc, "13.4.3", "Add file upload API for course materials (PDF/images)", badge="★ Advanced")
    task(doc, "13.4.4", "Add image upload for user profile avatar", badge="★ Advanced")

    section(doc, "13.5  Accessibility & Internationalisation")
    task(doc, "13.5.1", "Add screen reader support (Semantics widgets in Flutter)", badge="★ Advanced")
    task(doc, "13.5.2", "Support dynamic text size (MediaQuery.textScaleFactor)", badge="★ Advanced")
    task(doc, "13.5.3", "Add multi-language support with flutter_localizations (l10n)", badge="★ Advanced")
    task(doc, "13.5.4", "Add dark mode theme support", badge="★ Advanced")

    section(doc, "13.6  Admin & Analytics")
    task(doc, "13.6.1", "Build NGO admin web dashboard (React or Flutter Web)", badge="★ Advanced")
    task(doc, "13.6.2", "Add admin API: manage users, courses, counsellors", badge="★ Advanced")
    task(doc, "13.6.3", "Add learning analytics charts (weekly hours, mood trends)", badge="★ Advanced")
    task(doc, "13.6.4", "Add cohort-level reporting for NGO staff", badge="★ Advanced")

    section(doc, "13.7  Infrastructure Upgrades")
    task(doc, "13.7.1", "Migrate from SQLite to PostgreSQL for multi-user scale", badge="★ Advanced")
    task(doc, "13.7.2", "Add Redis cache for frequently-hit endpoints", badge="★ Advanced")
    task(doc, "13.7.3", "Add Alembic for database schema migrations", badge="★ Advanced")
    task(doc, "13.7.4", "Containerise backend with Docker + docker-compose", badge="★ Advanced")
    task(doc, "13.7.5", "Set up Kubernetes deployment for high-availability (production)", badge="★ Advanced")

    # ── PHASE 14 ──────────────────────────────────────────────────────────────
    phase_heading(doc, "PHASE 14", "User Management & Authentication", DARK_NAVY)
    section(doc, "14.1  User Registration & Login")
    task(doc, "14.1.1", "Design registration form: name, email/phone, password, role, age", important=True)
    task(doc, "14.1.2", "Implement POST /auth/register endpoint — hash password with bcrypt", important=True)
    note(doc, "Use passlib[bcrypt]: pwd_context.hash(password) — NEVER store plain-text passwords")
    task(doc, "14.1.3", "Implement POST /auth/login — verify password, return signed JWT access token", important=True)
    note(doc, "Use python-jose: jwt.encode({sub, role, exp}, SECRET_KEY, algorithm='HS256')")
    task(doc, "14.1.4", "Implement JWT token refresh endpoint (POST /auth/refresh)")
    task(doc, "14.1.5", "Add token expiry: access token 15 min, refresh token 7 days")
    task(doc, "14.1.6", "Implement password reset flow — send reset link via email (optional: SMS)")
    task(doc, "14.1.7", "Implement email verification on registration (optional)")
    task(doc, "14.1.8", "Add account lockout after 5 failed login attempts", important=True)
    task(doc, "14.1.9", "Store refresh tokens in DB — invalidate on logout (token revocation)")
    task(doc, "14.1.10", "Write Flutter login screen: email, password fields, Login button")
    task(doc, "14.1.11", "Write Flutter registration screen with role selector")
    task(doc, "14.1.12", "Store JWT access token securely using flutter_secure_storage", important=True)
    task(doc, "14.1.13", "Auto-attach Bearer token to all outgoing API requests in ApiClient")
    task(doc, "14.1.14", "Handle 401 Unauthorized — redirect to login screen automatically")

    section(doc, "14.2  User Profile Management")
    task(doc, "14.2.1", "Implement GET /users/me — return profile from JWT claims")
    task(doc, "14.2.2", "Implement PATCH /users/me — update name, bio, contact details")
    task(doc, "14.2.3", "Implement POST /users/me/avatar — upload profile photo")
    note(doc, "Save image to storage (local disk or S3); store URL in DB")
    task(doc, "14.2.4", "Implement POST /users/me/change-password — verify old password before update")
    task(doc, "14.2.5", "Implement DELETE /users/me — soft-delete (set is_active=False)")
    task(doc, "14.2.6", "Display profile photo in Flutter ProfileHeroCard")
    task(doc, "14.2.7", "Add editable fields in ProfileScreen (name, bio)")
    task(doc, "14.2.8", "Validate new password meets minimum complexity rules (8+ chars, 1 digit)")

    section(doc, "14.3  Multi-Role Account Types")
    task(doc, "14.3.1", "Add role field to User model (student / parent / mentor / admin / super_admin)")
    task(doc, "14.3.2", "Create parent accounts linked to child student accounts (parent_of FK)")
    task(doc, "14.3.3", "Implement Parent Dashboard with PIN protection (4-digit PIN stored hashed)")
    task(doc, "14.3.4", "Create mentor accounts — mentors linked to student cohorts")
    task(doc, "14.3.5", "Create NGO admin accounts — admin can manage all users and content")
    task(doc, "14.3.6", "Create super_admin role — full system access including role assignment")
    task(doc, "14.3.7", "Add is_active Boolean flag to all account types (default True)")
    task(doc, "14.3.8", "Add account approval workflow — new mentor/admin accounts need super_admin approval")
    task(doc, "14.3.9", "Test all role-based flows end-to-end with separate accounts")
    spacer(doc)

    # ── PHASE 15 ──────────────────────────────────────────────────────────────
    phase_heading(doc, "PHASE 15", "Role-Based Access Control (RBAC)", SOFT_RED)
    section(doc, "15.1  Role Definition & Permission Matrix")
    task(doc, "15.1.1", "Define all platform roles: Super Admin · Admin · Mentor · Content Creator · Student · Viewer/Guest")
    task(doc, "15.1.2", "Create a permission matrix document (rows = roles, cols = actions)")
    note(doc, "Actions: view_content · enroll_course · create_quiz · moderate_content · manage_users · view_analytics · run_drive")
    task(doc, "15.1.3", "Define which roles can access which API endpoints (annotate each router)")
    task(doc, "15.1.4", "Define feature-level access: e.g. only Mentors can create quizzes; only Admins can run drives")
    task(doc, "15.1.5", "Document the permission matrix in CONTRIBUTING.md and Swagger descriptions")

    section(doc, "15.2  Database Role Schema")
    task(doc, "15.2.1", "Add role column (Enum) to User table: super_admin/admin/mentor/content_creator/student/guest")
    task(doc, "15.2.2", "Create permissions table: id, name (e.g. 'manage_users'), description")
    task(doc, "15.2.3", "Create role_permissions junction table: role_name ↔ permission_name")
    task(doc, "15.2.4", "Seed permission matrix into DB at startup from config", important=True)
    task(doc, "15.2.5", "Add DB index on user.role for fast role-based queries")

    section(doc, "15.3  Backend RBAC Implementation")
    task(doc, "15.3.1", "Include role claim in JWT payload at login: {'sub': user_id, 'role': user.role}", important=True)
    task(doc, "15.3.2", "Create get_current_user dependency — decode JWT, fetch user from DB", important=True)
    task(doc, "15.3.3", "Create require_role(*roles) dependency factory — raises 403 if role not in allowed list", important=True)
    note(doc, "Usage: Depends(require_role('admin', 'super_admin')) in router")
    task(doc, "15.3.4", "Apply role guards to all sensitive endpoints", important=True)
    note(doc, "e.g. POST /courses — Content Creator+; DELETE /users/{id} — Admin+; GET /analytics — Mentor+")
    task(doc, "15.3.5", "Add RBAC middleware that logs all access attempts with user_id + role + endpoint")
    task(doc, "15.3.6", "Implement token/session handling — invalidate tokens on role change", important=True)
    task(doc, "15.3.7", "Add POST /auth/logout — blacklist refresh token in DB")
    task(doc, "15.3.8", "Ensure users can only access their own data (ownership check on top of role check)", important=True)
    note(doc, "e.g. Student can GET /users/{id}/courses only if id == their own user_id")

    section(doc, "15.4  Frontend Route Protection")
    task(doc, "15.4.1", "Parse role from JWT claims after login; store in AppState.role")
    task(doc, "15.4.2", "Create RoleGuard widget — shows child only if user has required role")
    task(doc, "15.4.3", "Hide admin navigation tabs for Student/Guest roles")
    task(doc, "15.4.4", "Protect admin screens with route-level role check — redirect to home if unauthorized")
    task(doc, "15.4.5", "Show/hide quiz-creation button based on Mentor/ContentCreator role")
    task(doc, "15.4.6", "Show/hide drive management controls based on Admin role")
    task(doc, "15.4.7", "Display 'Access Denied' screen for unauthorized navigation attempts")

    section(doc, "15.5  RBAC Testing")
    task(doc, "15.5.1", "Write test: Student cannot access POST /courses (expect 403)")
    task(doc, "15.5.2", "Write test: Mentor can create quiz, but cannot delete other users (403)")
    task(doc, "15.5.3", "Write test: Admin can manage all users")
    task(doc, "15.5.4", "Write test: Student cannot view another student's mood logs (403)")
    task(doc, "15.5.5", "Write test: Guest can view public course catalogue but not enroll (401)")
    task(doc, "15.5.6", "Write test: Expired JWT returns 401 on all protected endpoints")
    task(doc, "15.5.7", "Write test: Tampered JWT signature returns 401", important=True)
    task(doc, "15.5.8", "Manual test: log in as each role, verify correct UI is shown in Flutter")
    spacer(doc)

    # ── PHASE 16 ──────────────────────────────────────────────────────────────
    phase_heading(doc, "PHASE 16", "Quiz & Daily Challenge System", ORANGE)
    section(doc, "16.1  Database Schema — Quiz Engine")
    task(doc, "16.1.1", "Create quizzes table: id, title, description, category_id, created_by, status (draft/published/archived), scheduled_at")
    task(doc, "16.1.2", "Create questions table: id, quiz_id, question_text, question_type (mcq/true_false/fill), points, order_index")
    task(doc, "16.1.3", "Create options table: id, question_id, option_text, is_correct")
    task(doc, "16.1.4", "Create quiz_attempts table: id, user_id, quiz_id, score, max_score, completed_at, time_taken_seconds")
    task(doc, "16.1.5", "Create quiz_answers table: id, attempt_id, question_id, selected_option_id, is_correct")
    task(doc, "16.1.6", "Create daily_challenges table: id, quiz_id, challenge_date (unique), xp_bonus")
    task(doc, "16.1.7", "Create quiz_categories table: id, name, color_hex, icon_name")
    task(doc, "16.1.8", "Add foreign keys and cascade rules for all quiz tables")
    task(doc, "16.1.9", "Add indexes: quiz_attempts(user_id), daily_challenges(challenge_date)")

    section(doc, "16.2  Backend APIs — Quiz CRUD")
    task(doc, "16.2.1", "POST /quizzes — Mentor/ContentCreator creates new quiz (status=draft)")
    task(doc, "16.2.2", "GET /quizzes — list published quizzes; Admin sees all including drafts")
    task(doc, "16.2.3", "GET /quizzes/{id} — get quiz with all questions and options")
    task(doc, "16.2.4", "PUT /quizzes/{id} — edit quiz (owner or Admin only)")
    task(doc, "16.2.5", "DELETE /quizzes/{id} — soft-delete (set status=archived, Admin only)")
    task(doc, "16.2.6", "POST /quizzes/{id}/publish — Admin approves and publishes quiz")
    task(doc, "16.2.7", "POST /quizzes/upload — bulk upload quiz from JSON/CSV file (Mentor+)")
    note(doc, "Parse JSON file: [{title, questions: [{text, options: [{text, correct}]}]}]")
    task(doc, "16.2.8", "POST /quizzes/{id}/questions — add a question to an existing quiz")
    task(doc, "16.2.9", "GET /quiz-categories — list all quiz categories")

    section(doc, "16.3  Backend APIs — Daily Challenge & Attempts")
    task(doc, "16.3.1", "GET /daily-challenge — return today's challenge (lookup by challenge_date = today)")
    task(doc, "16.3.2", "POST /daily-challenge — Admin schedules a quiz as daily challenge for a date")
    task(doc, "16.3.3", "POST /quizzes/{id}/attempt — Student starts a quiz attempt (create attempt record)")
    task(doc, "16.3.4", "POST /quizzes/{id}/attempt/{attempt_id}/submit — submit answers, calculate score")
    note(doc, "For each answer: match selected_option_id to correct option; sum points; store results")
    task(doc, "16.3.5", "POST /quizzes/{id}/attempt after submit — award XP via user_crud.add_xp")
    task(doc, "16.3.6", "GET /users/{id}/quiz-history — list all attempts for a student")
    task(doc, "16.3.7", "GET /quizzes/{id}/leaderboard — top scores for a specific quiz")
    task(doc, "16.3.8", "GET /quizzes/analytics/{id} — Admin quiz report: attempts, avg score, completion rate")

    section(doc, "16.4  Quiz Frontend — Student View")
    task(doc, "16.4.1", "Build QuizListScreen — show published quizzes with category filter")
    task(doc, "16.4.2", "Build DailyChallengeCard — show today's challenge on Home screen")
    task(doc, "16.4.3", "Build QuizPlayerScreen — show question, options, timer, progress bar")
    task(doc, "16.4.4", "Build ResultScreen — show score, XP earned, correct answers breakdown")
    task(doc, "16.4.5", "Build QuizLeaderboardScreen — show top 10 scores for current quiz")
    task(doc, "16.4.6", "Display 'Quiz already completed today' if daily challenge already attempted")
    task(doc, "16.4.7", "Show badge/XP earned animation on quiz completion (gamification)")

    section(doc, "16.5  Quiz Frontend — Mentor/Admin View")
    task(doc, "16.5.1", "Build QuizCreatorScreen — form to create quiz with questions and options")
    task(doc, "16.5.2", "Build QuizUploadScreen — file picker to upload JSON/CSV quiz file")
    task(doc, "16.5.3", "Build DailyChallengeSchedulerScreen — Admin sets daily challenge calendar")
    task(doc, "16.5.4", "Build QuizModerationScreen — Admin sees pending quizzes, approve/reject")
    task(doc, "16.5.5", "Build QuizAnalyticsScreen — charts for completion rate, avg score, student performance")
    task(doc, "16.5.6", "Add gamification badges auto-award: 'Quiz Master' after 10 quizzes completed")
    task(doc, "16.5.7", "Add streak tracking: 'Daily Streak' badge for 7 consecutive daily challenges")
    spacer(doc)

    # ── PHASE 17 ──────────────────────────────────────────────────────────────
    phase_heading(doc, "PHASE 17", "Tarase Heerah Drive — Talent Selection", RGBColor(0x70, 0xD9, 0x8B))
    section(doc, "17.1  Module Overview & Design")
    task(doc, "17.1.1", "Define the Tarase Heerah Drive: periodic talent selection event for deserving students")
    note(doc, "Students participate in challenges/quizzes → Admin shortlists → Lucky draw OR manual selection → Counselling access")
    task(doc, "17.1.2", "Define participation criteria (minimum quiz score, attendance, age range)")
    task(doc, "17.1.3", "Define selection outcomes: counselling session access, mentorship, rewards")
    task(doc, "17.1.4", "Design admin workflow: create drive → open participation → shortlist → select → notify → schedule counselling")

    section(doc, "17.2  Database Schema — Drive Engine")
    task(doc, "17.2.1", "Create talent_drives table: id, name, description, status (draft/open/shortlisting/closed), start_date, end_date, selection_method (lucky_draw/manual/hybrid), max_selections")
    task(doc, "17.2.2", "Create drive_eligibility_criteria table: id, drive_id, min_quiz_score, min_quizzes_completed, age_min, age_max")
    task(doc, "17.2.3", "Create drive_participants table: id, drive_id, user_id, total_score, quizzes_completed, is_shortlisted, status (pending/shortlisted/selected/not_selected)")
    task(doc, "17.2.4", "Create drive_selections table: id, drive_id, user_id, selected_at, selected_by (admin_user_id), selection_note, counselling_session_id")
    task(doc, "17.2.5", "Create drive_notifications table: id, drive_id, user_id, message, notification_type (selected/shortlisted/not_selected), sent_at, is_read")
    task(doc, "17.2.6", "Add indexes: drive_participants(drive_id, user_id), drive_selections(drive_id)")

    section(doc, "17.3  Backend APIs — Drive Management")
    task(doc, "17.3.1", "POST /drives — Admin creates a new Tarase Heerah Drive")
    task(doc, "17.3.2", "GET /drives — list all drives; students see open drives; admins see all")
    task(doc, "17.3.3", "GET /drives/{id} — drive details including eligibility criteria")
    task(doc, "17.3.4", "PATCH /drives/{id}/status — Admin opens/closes/moves drive through stages")
    task(doc, "17.3.5", "POST /drives/{id}/participate — Student registers participation")
    task(doc, "17.3.6", "GET /drives/{id}/participants — Admin views all participants with scores")
    task(doc, "17.3.7", "POST /drives/{id}/shortlist — Admin marks students as shortlisted (batch operation)")
    task(doc, "17.3.8", "POST /drives/{id}/select/lucky-draw — System randomly selects N students from shortlisted pool")
    note(doc, "Use Python random.sample(shortlisted_ids, k=max_selections) — seed with drive_id for reproducibility")
    task(doc, "17.3.9", "POST /drives/{id}/select/manual — Admin manually selects specific student IDs")
    task(doc, "17.3.10", "POST /drives/{id}/select/hybrid — Lucky draw from shortlist, Admin can override/replace any selection")
    task(doc, "17.3.11", "GET /drives/{id}/selections — view final selected students")
    task(doc, "17.3.12", "GET /drives/{id}/history — full audit trail of selection actions")

    section(doc, "17.4  Backend APIs — Counselling & Notifications")
    task(doc, "17.4.1", "POST /drives/{id}/notify — send in-app notifications to all participants with outcome")
    task(doc, "17.4.2", "POST /drives/{id}/schedule-counselling — auto-create counselling sessions for all selected students")
    task(doc, "17.4.3", "GET /users/{id}/drive-history — student's participation and selection history across all drives")
    task(doc, "17.4.4", "GET /users/{id}/notifications — all unread notifications for a user")
    task(doc, "17.4.5", "PATCH /notifications/{id}/read — mark notification as read")
    task(doc, "17.4.6", "GET /drives/{id}/counselling-sessions — list all counselling sessions created for this drive")

    section(doc, "17.5  Frontend — Student View")
    task(doc, "17.5.1", "Build DriveListScreen — show open drives with eligibility status indicator")
    task(doc, "17.5.2", "Build DriveDetailScreen — show drive description, criteria, prizes, participate button")
    task(doc, "17.5.3", "Build DriveResultScreen — show 'Congratulations! You have been selected' or 'Thank you for participating'")
    task(doc, "17.5.4", "Build NotificationsScreen — list all drive notifications with read/unread state")
    task(doc, "17.5.5", "Show badge on notification bell icon when unread notifications exist")
    task(doc, "17.5.6", "Display assigned counselling session in WellnessScreen after selection")

    section(doc, "17.6  Frontend — Admin View")
    task(doc, "17.6.1", "Build DriveManagerScreen — create/edit drives with eligibility criteria form")
    task(doc, "17.6.2", "Build ParticipantListScreen — data table of participants, scores, shortlist checkboxes")
    task(doc, "17.6.3", "Build SelectionScreen — choose selection method (lucky draw / manual / hybrid), run selection")
    task(doc, "17.6.4", "Show selection results with student names, scores, and counselling status")
    task(doc, "17.6.5", "Build drive history/audit log view for admin review")
    task(doc, "17.6.6", "Test end-to-end: create drive → student participates in quiz → shortlist → lucky draw → notify → counselling scheduled")
    spacer(doc)

    # ── PHASE 18 ──────────────────────────────────────────────────────────────
    phase_heading(doc, "PHASE 18", "Learning Content Platform", BLUE)
    section(doc, "18.1  Database Schema — Content")
    task(doc, "18.1.1", "Extend courses table: is_free (Boolean), price (Decimal), thumbnail_url, trailer_video_url, published_at, content_status (draft/review/published/archived)")
    task(doc, "18.1.2", "Create course_modules table: id, course_id, title, order_index, is_free_preview")
    task(doc, "18.1.3", "Create lessons table: id, module_id, title, lesson_type (video/pdf/text/quiz), duration_seconds, order_index, content_url")
    task(doc, "18.1.4", "Create user_lesson_progress table: id, user_id, lesson_id, watched_seconds, completed, completed_at")
    task(doc, "18.1.5", "Create course_enrollments table: id, user_id, course_id, enrolled_at, payment_status (free/paid/pending)")
    task(doc, "18.1.6", "Create content_reports table: id, user_id, content_id, content_type, reason, status (pending/reviewed/resolved)")
    task(doc, "18.1.7", "Create cybersecurity_content table: id, title, category (phishing/data_privacy/password_safety/online_bullying), difficulty, content_md")
    task(doc, "18.1.8", "Add indexes: user_lesson_progress(user_id, lesson_id), course_enrollments(user_id)")

    section(doc, "18.2  Backend APIs — Course Content Management")
    task(doc, "18.2.1", "POST /courses — ContentCreator/Admin creates course with is_free flag")
    task(doc, "18.2.2", "POST /courses/{id}/modules — add module to course")
    task(doc, "18.2.3", "POST /modules/{id}/lessons — add lesson (video/pdf/text) to module")
    task(doc, "18.2.4", "POST /lessons/upload-video — upload video file; return storage URL")
    note(doc, "Store locally under media/videos/ in dev; use S3 or Cloudinary in production")
    task(doc, "18.2.5", "POST /lessons/upload-document — upload PDF/image; return URL")
    task(doc, "18.2.6", "PATCH /courses/{id}/status — Admin publishes/archives course (content moderation workflow)")
    task(doc, "18.2.7", "POST /courses/{id}/enroll — Student enrolls in free course (or initiates payment for paid)")
    task(doc, "18.2.8", "GET /courses/{id}/curriculum — full course structure (modules + lessons)")
    task(doc, "18.2.9", "PATCH /users/{id}/lessons/{lid}/progress — update watched_seconds and completed")
    task(doc, "18.2.10", "GET /users/{id}/courses/{cid}/progress — detailed lesson-by-lesson progress")
    task(doc, "18.2.11", "POST /content-reports — Student/Parent reports inappropriate content", important=True)
    task(doc, "18.2.12", "GET /admin/content-reports — Admin views and resolves pending reports")

    section(doc, "18.3  Cybersecurity Awareness Content")
    task(doc, "18.3.1", "Seed cybersecurity lessons: Phishing Awareness, Safe Passwords, Data Privacy, Cyberbullying, Screen Time")
    task(doc, "18.3.2", "Add POST /cybersecurity — Admin/ContentCreator creates new awareness article")
    task(doc, "18.3.3", "Add GET /cybersecurity?category=phishing — filter by category")
    task(doc, "18.3.4", "Link cybersecurity content to quiz system: attach a quiz at end of each lesson")
    task(doc, "18.3.5", "Add InternetSafetyScreen in Flutter with cybersecurity lesson cards")
    task(doc, "18.3.6", "Award 'Cyber Safe' badge after completing all cybersecurity modules")

    section(doc, "18.4  Video Streaming Optimisation")
    task(doc, "18.4.1", "Integrate video_player Flutter package for in-app video playback")
    task(doc, "18.4.2", "Show video progress bar and resume from last watched position")
    task(doc, "18.4.3", "Mark lesson as completed when video watched > 90%")
    task(doc, "18.4.4", "Add video quality selector (360p / 720p) for low-bandwidth users")
    task(doc, "18.4.5", "Cache video metadata (not content) for offline catalogue browsing")
    task(doc, "18.4.6", "Implement adaptive bitrate streaming (HLS) for production (★ Advanced)")
    note(doc, "Use FFmpeg server-side to transcode uploads to HLS format (.m3u8 playlist + .ts segments)")

    section(doc, "18.5  Content Moderation Workflow")
    task(doc, "18.5.1", "New course submitted by ContentCreator → status = 'review'")
    task(doc, "18.5.2", "Admin receives in-app notification: 'New course pending review'")
    task(doc, "18.5.3", "Admin reviews course content: preview video, read lessons, check quiz")
    task(doc, "18.5.4", "Admin approves → status = 'published'; student feed updates immediately")
    task(doc, "18.5.5", "Admin rejects → status = 'draft', rejection reason sent to ContentCreator")
    task(doc, "18.5.6", "Student reports content → admin receives report → admin reviews and removes if needed", important=True)
    task(doc, "18.5.7", "Log all moderation actions with admin_id, timestamp, action, reason")
    task(doc, "18.5.8", "Verify child-safe content standards: no violence, no adult content, age-appropriate language", important=True)

    section(doc, "18.6  Progress Tracking & Certificates")
    task(doc, "18.6.1", "Calculate course completion %: completed_lessons / total_lessons × 100")
    task(doc, "18.6.2", "Mark course as completed when all lessons done and final quiz passed")
    task(doc, "18.6.3", "Auto-award XP on lesson completion (XP = lesson duration ÷ 60 minutes × 100)")
    task(doc, "18.6.4", "Generate completion certificate PDF using reportlab or weasyprint (★ Advanced)")
    note(doc, "Certificate includes student name, course name, completion date, NGO logo")
    task(doc, "18.6.5", "Show course completion badge on ProfileScreen")
    task(doc, "18.6.6", "Build LearningDashboard: weekly learning hours chart, streak, completed courses")
    task(doc, "18.6.7", "Add skill-path feature: recommended course sequence per skill category")
    spacer(doc)

    # Footer note
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        f"CareSkill Project Checklist  ·  Generated {datetime.date.today().strftime('%d %B %Y')}  ·  v1.0.0"
    )
    fr.italic          = True
    fr.font.size       = Pt(10)
    fr.font.color.rgb  = MUTED


# ── MAIN ───────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Page margins (A4-friendly)
    for section_obj in doc.sections:
        section_obj.top_margin    = Cm(2.0)
        section_obj.bottom_margin = Cm(2.0)
        section_obj.left_margin   = Cm(2.5)
        section_obj.right_margin  = Cm(2.5)

    # Default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Heading 1 style
    h1 = doc.styles["Heading 1"]
    h1.font.name       = "Calibri"
    h1.font.size       = Pt(16)
    h1.font.bold       = True
    h1.font.color.rgb  = DARK_NAVY

    add_title_page(doc)
    add_toc(doc)
    build_checklist(doc)

    output = "project_checklist.docx"
    doc.save(output)
    print(f"✓  Saved: {output}")
    print(f"   Phases:  18")
    print(f"   Tasks:   ~280 items across the full SDLC")


if __name__ == "__main__":
    main()
